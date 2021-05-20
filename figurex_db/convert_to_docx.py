"""
Usage:
    script.py [options] -f figure_folder -s text_json_file -d docx_file

Options:
    -f <directory>      Figure folder
    -s <file>           Text json file
    -d <file>           Docx file
"""

import collections
import json
from pathlib import Path

import docopt
import tqdm
from docx import Document
from docx.shared import Inches

from figurex_db.utils import generate_path


def to_docx(bioc_dir, src, dest):
    bioc_dir = Path(bioc_dir)

    with open(src) as fp:
        objs = json.load(fp)

    document = Document()

    cnt = collections.Counter()
    for i, doc in tqdm.tqdm(enumerate(objs)):
        pmcid = doc['pmcid']
        figure_name = doc['url']
        figure_name = figure_name[figure_name.rfind('/') + 1:]
        # get figure
        figure_file = bioc_dir / generate_path(pmcid) / '{}_{}'.format(pmcid, figure_name)
        # get subfigure
        for subfigure in doc['files']:
            if subfigure['type'] in ['ct', 'cxr']:
                if subfigure['is_subfigure']:
                    xtl = subfigure['xtl']
                    ytl = subfigure['ytl']
                    xbr = subfigure['xbr']
                    ybr = subfigure['ybr']
                    local_file = figure_file.parent / f'{figure_file.stem}_{xtl}x{ytl}_{xbr}x{ybr}{figure_file.suffix}'
                else:
                    local_file = figure_file
                try:
                    document.add_picture(str(local_file), width=Inches(3))
                    document.add_paragraph('Type: {}, PMC: {}'.format(subfigure['type'], doc['pmcid']))
                except:
                    print('Cannot add', local_file)
                    continue
        # get caption
        document.add_heading('Caption', level=1)
        document.add_paragraph(doc['caption']['text'])

        # get referred text
        document.add_heading('Referred text', level=1)
        for r in doc['referred_text']:
            document.add_paragraph(r['text'], style='List Bullet')

        if i >= 10:
            break

    document.save(str(dest))
    print(cnt)


if __name__ == '__main__':
    args = docopt.docopt(__doc__)
    to_docx(bioc_dir=args['-f'], src=args['-s'], dest=args['-d'])